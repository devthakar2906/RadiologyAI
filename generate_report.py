from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUTPUT_PATH = r"E:\Dev Microsoft\ChatGPT\RadiologyAI\Clg\RadiologyAI_Final_Project_Report.docx"
INSTITUTE_NAME = "SAL INSTITUTE OF TECHNOLOGY AND ENGINEERING RESEARCH"
INSTITUTE_ADDRESS = "App. Science City, Sola Road, Ahmedabad - 380060"
PROJECT_ID = "230673107042"


def set_page_layout(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)


def set_default_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_header_footer(section, chapter_heading: str) -> None:
    header = section.header
    if not header.paragraphs:
        header.add_paragraph()
    hp = header.paragraphs[0]
    hp.text = f"{PROJECT_ID}\t\t{chapter_heading}"
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.add_run("Gujarat Technological University")
    fp.add_run("\t")
    add_page_number(fp)
    fp.add_run("\t")
    fp.add_run(INSTITUTE_NAME)


def add_blank(document: Document, count: int = 1):
    for _ in range(count):
        document.add_paragraph()


def add_centered(document: Document, text: str, size: int = 14, bold: bool = False, italic: bool = False):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic


def add_heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text.upper() if level <= 2 else text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    r.bold = True


def add_para(document: Document, text: str, italic: bool = False):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    r.italic = italic


def add_toc_placeholder(document: Document):
    add_heading(document, "Table of Contents", 1)
    add_para(document, "To update the table of contents in Word, click in the TOC and press F9.", italic=True)
    p = document.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and update field."
    fld_sep.append(t)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_sep)
    p._p.append(fld_end)


def add_cover_pages(doc: Document):
    add_blank(doc, 4)
    for line in ["GUJARAT TECHNOLOGICAL UNIVERSITY", "Ahmedabad"]:
        add_centered(doc, line, 16, True)
    add_blank(doc, 2)
    add_centered(doc, "RADIOLOGY AI PLATFORM", 18, True)
    add_blank(doc, 1)
    add_centered(doc, "A", 14, True)
    add_centered(doc, "PROJECT REPORT", 14, True)
    add_blank(doc, 1)
    add_centered(doc, "Submitted by", 14, italic=True)
    add_blank(doc, 1)
    add_centered(doc, "DEV THAKAR", 16, True)
    add_centered(doc, "230673107042", 14)
    add_blank(doc, 1)
    add_centered(doc, "In partial fulfillment for the award of the degree of", 14, italic=True)
    add_centered(doc, "BACHELOR OF ENGINEERING", 16, True)
    add_centered(doc, "in", 14)
    add_centered(doc, "COMPUTER ENGINEERING", 14, True)
    add_blank(doc, 1)
    add_centered(doc, INSTITUTE_NAME, 14)
    add_centered(doc, INSTITUTE_ADDRESS, 14)
    add_blank(doc, 1)
    add_centered(doc, "Gujarat Technological University, Ahmedabad", 16)
    add_centered(doc, "April, 2026", 14)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_blank(doc, 3)
    for line in ["GUJARAT TECHNOLOGICAL UNIVERSITY", "Ahmedabad"]:
        add_centered(doc, line, 16, True)
    add_blank(doc, 2)
    add_centered(doc, "RADIOLOGY AI PLATFORM", 18, True)
    add_blank(doc, 1)
    add_centered(doc, "A", 14, True)
    add_centered(doc, "PROJECT REPORT", 14, True)
    add_blank(doc, 1)
    add_centered(doc, "Submitted by", 14, italic=True)
    add_blank(doc, 1)
    add_centered(doc, "DEV THAKAR", 16, True)
    add_centered(doc, "(230673107042)", 14)
    add_blank(doc, 1)
    add_centered(doc, "In partial fulfillment for the award of the degree of", 14, italic=True)
    add_centered(doc, "BACHELOR OF ENGINEERING", 16, True)
    add_centered(doc, "IN", 14, True)
    add_centered(doc, "COMPUTER ENGINEERING", 14, True)
    add_blank(doc, 1)
    add_centered(doc, INSTITUTE_NAME, 14)
    add_centered(doc, INSTITUTE_ADDRESS, 14)
    add_blank(doc, 1)
    add_centered(doc, "Gujarat Technological University, Ahmedabad", 16)
    add_centered(doc, "Academic Year: 2025-26", 14)


def add_prefatory_sections(doc: Document):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_centered(doc, INSTITUTE_NAME, 14)
    add_centered(doc, INSTITUTE_ADDRESS, 14)
    add_blank(doc, 1)
    add_heading(doc, "Certificate", 1)
    add_para(doc, 'This is to certify that the project report submitted along with the project entitled "Radiology AI Platform" has been carried out by Dev Thakar under my guidance in partial fulfillment for the degree of Bachelor of Engineering in Computer Engineering, 8th Semester of Gujarat Technological University, Ahmedabad during the academic year 2025-26.')
    add_blank(doc, 2)
    for line in [
        "_________________________              _________________________",
        "Komal Thummar                          Dr. Nimisha Patel",
        "Internal Guide                         Head of the Department",
    ]:
        add_para(doc, line)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    for line in [
        "FICE Education Pvt. Ltd.",
        "Remote Internship",
        "Date: 01/04/2026",
    ]:
        add_para(doc, line)
    add_blank(doc, 1)
    add_heading(doc, "To Whom It May Concern", 1)
    for para in [
        "This is to certify that Dev Thakar, a student of Sal Institute of Technology and Engineering Research, has successfully completed his internship in the field of Full Stack Software Development and AI-based Reporting Systems from 21/01/2026 to 01/04/2026 (Total number of Weeks: 12) under the guidance of Ms. Anjali Singh.",
        "His internship activities include requirement analysis, backend API development, frontend dashboard implementation, asynchronous speech-processing workflow, structured report generation, encrypted data handling, database integration, and deployment-oriented system testing for the Radiology AI Platform.",
        "During the period of his internship program with us, he was exposed to different software engineering processes and was found diligent, hardworking, and inquisitive.",
        "We wish him every success in his life and career.",
    ]:
        add_para(doc, para)
    add_blank(doc, 2)
    add_para(doc, "For FICE Education Pvt. Ltd.")
    add_blank(doc, 2)
    add_para(doc, "_________________________")
    add_para(doc, "Anjali Singh")
    add_para(doc, "Industry Mentor")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_centered(doc, INSTITUTE_NAME, 14)
    add_centered(doc, INSTITUTE_ADDRESS, 14)
    add_blank(doc, 1)
    add_heading(doc, "Declaration", 1)
    add_para(doc, 'I hereby declare that the Internship / Project report submitted along with the Internship / Project entitled "Radiology AI Platform" submitted in partial fulfillment for the degree of Bachelor of Engineering in Computer Engineering to Gujarat Technological University, Ahmedabad, is a bonafide record of original project work carried out by me at FICE Education Pvt. Ltd. under the supervision of Ms. Anjali Singh and Ms. Komal Thummar and that no part of this report has been directly copied from any students\' reports or taken from any other source, without providing due reference.')
    add_blank(doc, 2)
    add_para(doc, "Name of Student                         Sign of Student")
    add_para(doc, "Dev Thakar                              _________________________")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "Acknowledgement", 1)
    for para in [
        "It gives me immense pleasure to present this project report on “Radiology AI Platform”. I would like to express my sincere gratitude to all those who have guided, supported, and encouraged me throughout the development of this project.",
        "First and foremost, I am highly thankful to my internal guide, Ms. Komal Thummar, for her valuable guidance, continuous encouragement, and constructive suggestions, which greatly helped me in planning, designing, and completing this work in a systematic manner.",
        "I am equally grateful to Dr. Nimisha Patel, Head of the Department of Computer Engineering, for providing the necessary academic environment, infrastructure, and motivation required for the successful completion of this project.",
        "I would also like to thank my industry mentor, Ms. Anjali Singh, and FICE Education Pvt. Ltd. for providing me with the opportunity to work on a practical and meaningful industry-oriented project. Their support, guidance, and exposure to real-world requirements helped me understand how software systems are designed and deployed to solve domain-specific problems.",
        "I extend my heartfelt thanks to the faculty members of the Computer Engineering Department and my institute for their support and encouragement during the project period.",
        "Lastly, I am deeply thankful to my family and friends for their constant motivation, patience, and moral support throughout this journey.",
    ]:
        add_para(doc, para)
    add_blank(doc, 1)
    add_para(doc, "DEV THAKAR")
    add_para(doc, "230673107042")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "Abstract", 1)
    for para in [
        "Radiology reporting is a critical part of clinical workflow, but conventional dictation and manual report preparation can be time-consuming, repetitive, and error-prone. Doctors often need to convert free-form radiology findings into standardized, structured reports while managing high reporting volume. The project “Radiology AI Platform” addresses this challenge by providing a full-stack, AI-enabled dictation and structured reporting system designed to improve reporting speed, consistency, and accessibility.",
        "The proposed system combines modern web technologies with AI-driven processing to enable live dictation, transcript refinement, structured report generation, secure storage, search functionality, and exportable report output. The frontend is built using React and Tailwind CSS to provide a responsive and user-friendly interface. The backend is developed using FastAPI and integrates PostgreSQL for persistent storage, Redis for caching and messaging, and Celery for asynchronous processing. Hugging Face APIs are used for transcription refinement and structured report assistance, while template-based formatting helps produce organized radiology reports.",
        "The platform supports role-based access for doctors and administrators, live browser speech capture, audio upload, asynchronous processing of transcription and report tasks, encrypted storage, report retrieval, filtering, search, and PDF export. The system has been designed to be modular, scalable, and practical for real-world use in radiology-oriented workflows.",
        "This project demonstrates how AI-assisted systems can improve medical documentation by reducing repetitive effort, maintaining structure, and supporting clinicians with a more efficient digital reporting pipeline.",
    ]:
        add_para(doc, para, italic=True)


def add_lists_and_toc(doc: Document):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "List of Figures", 1)
    for line in [
        "Fig 3.1 Overall System Architecture",
        "Fig 4.1 Authentication and Access Flow",
        "Fig 4.2 Live Dictation and Transcript Refinement Flow",
        "Fig 4.3 Structured Report Generation Workflow",
        "Fig 4.4 Celery and Redis Background Processing Flow",
        "Fig 5.1 Homepage / Landing Interface",
        "Fig 5.2 Reporting Dashboard Interface",
        "Fig 5.3 Structured Report Panel and Save Workflow",
        "Fig 6.1 Generated Report Output",
        "Fig 6.2 PDF Export Example",
    ]:
        add_para(doc, line)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "List of Tables", 1)
    for line in [
        "Table 2.1 Literature Review Comparison",
        "Table 3.1 Functional Modules of the Proposed System",
        "Table 3.2 Technology Stack Used",
        "Table 4.1 API and Backend Responsibilities",
        "Table 5.1 Key Frontend Components",
        "Table 6.1 Test Scenarios and Results",
        "Table 6.2 Performance and Reliability Observations",
    ]:
        add_para(doc, line)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "List of Abbreviations", 1)
    for line in [
        "AI      Artificial Intelligence",
        "API     Application Programming Interface",
        "ASR     Automatic Speech Recognition",
        "AES     Advanced Encryption Standard",
        "DB      Database",
        "JWT     JSON Web Token",
        "LLM     Large Language Model",
        "PDF     Portable Document Format",
        "UI      User Interface",
        "UX      User Experience",
        "SQL     Structured Query Language",
    ]:
        add_para(doc, line)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_toc_placeholder(doc)


CHAPTERS = [
    ("Chapter 1", "Introduction", [
        ("1.1 Background", [
            "Radiology is a data-intensive medical discipline in which accurate interpretation and clear communication are essential. After examining an imaging study, radiologists or clinicians dictate findings and convert them into formal reports for diagnosis, treatment planning, and record maintenance. In many settings, this process still depends on manual dictation and narrative report writing, which increases repetition, consumes time, and leads to inconsistency in structure and terminology.",
            "With the growth of artificial intelligence, speech technology, and medical information systems, there is an opportunity to improve the radiology reporting workflow. An ideal system should assist doctors in capturing findings, refining speech into readable text, converting raw content into structured sections, and securely storing reports for future access. Such a system should also support concurrency, role-based access, search, and export for practical institutional use.",
            "The Radiology AI Platform has been developed to address these requirements. It is a full-stack application that combines browser-based dictation, backend AI processing, template-driven structured reporting, encrypted storage, and asynchronous task processing to support real-time and scalable report generation.",
        ]),
        ("1.2 Problem Statement", [
            "Traditional radiology reporting workflows face several limitations. Free-form dictation often produces unstructured and inconsistent reports. Manual editing and formatting consume significant clinician time. Repetitive documentation reduces productivity in high-volume environments. Unstructured reports make search, reuse, and standardization difficult. Heavy local AI models are difficult to deploy and scale in lightweight systems. Concurrent report processing requires background task handling and queue management to avoid blocking the main application.",
            "Therefore, a need exists for a lightweight yet scalable system that can convert radiology findings into structured reports using AI assistance while maintaining security, performance, and usability.",
        ]),
        ("1.3 Objectives", [
            "The major objectives of this project are to design a web-based radiology reporting platform with doctor and admin roles; capture radiology findings through live dictation and audio upload; refine transcription using AI-based processing through external APIs; generate structured reports using template-driven workflows; store reports securely using encrypted storage techniques; support asynchronous processing using Celery and Redis; provide search, filtering, report history, and PDF export; and create a practical and extensible system suitable for academic and real-world use.",
        ]),
        ("1.4 Scope", [
            "The scope of the project includes role-based login and signup for doctor and admin users, browser-based live dictation for immediate text capture, audio upload and transcription refinement workflow, template-based structured report generation from findings, secure storage of reports and transcription data, retrieval, editing, saving, searching, and PDF downloading of reports, and background processing for multiple concurrent reporting requests.",
            "The project does not replace the clinical judgment of a radiologist. Instead, it acts as an AI-assisted documentation platform to improve workflow efficiency and standardization.",
        ]),
        ("1.5 Significance of the Project", [
            "The significance of the proposed project lies in its ability to bridge the gap between speech-based data entry and structured medical documentation. It improves the reporting experience by reducing manual formatting, increasing consistency, and organizing outputs into usable structured formats. The project also demonstrates the practical integration of modern AI services, backend orchestration, frontend usability, and secure data handling in a healthcare-oriented software system.",
        ]),
    ]),
    ("Chapter 2", "Literature Survey", [
        ("2.1 Existing Radiology Reporting Practices", [
            "Conventional radiology reporting is usually performed through manual typing or voice dictation followed by editing. Although voice dictation speeds up data entry, the final report often remains narrative and inconsistent. Structured reporting systems have been recommended in radiology because they improve completeness, standardization, and readability. However, traditional structured reporting tools may be rigid and time-consuming if they require manual template filling.",
        ]),
        ("2.2 AI in Medical Speech and Documentation", [
            "Recent developments in AI have enabled automatic speech recognition, text refinement, and report generation. Speech models can convert clinical audio into text, while language models can assist in summarization and formatting. In medical domains, especially radiology, the challenge is not only transcription accuracy but also preserving clinical meaning, anatomical references, measurements, and diagnostic intent.",
        ]),
        ("2.3 Comparative Analysis of Existing Approaches", [
            "Existing solutions in the broad domain typically offer basic dictation-to-text systems, hospital-specific reporting software, template libraries for structured reports, or generic AI text generation tools. However, many of these approaches have limitations. Generic speech-to-text systems are not optimized for medical phrasing, many reporting systems are expensive or institution-specific, fully local AI models can require high hardware resources, unstructured AI generation may produce outputs that are difficult to trust or reuse, and some systems lack secure storage, background processing, or scalable architecture.",
            "The Radiology AI Platform improves upon these limitations by combining a lightweight web interface, API-based AI inference, template-guided structuring, encrypted storage, and scalable asynchronous processing.",
        ]),
        ("2.4 Research Gap", [
            "The key gap identified from existing practices is the lack of an integrated, lightweight, template-driven, and scalable reporting platform that supports live dictation and audio upload, refines radiology transcription through API-based AI, produces structured report output instead of generic prose, stores reports securely with role-based access, and handles multiple requests concurrently through background workers. The proposed project addresses this gap through a modern full-stack implementation focused on radiology documentation workflow improvement.",
        ]),
    ]),
    ("Chapter 3", "Methodology", [
        ("3.1 Proposed System Overview", [
            "The proposed system is a web-based radiology reporting platform designed around a modular architecture. It accepts radiology findings from typed text, live dictation, or uploaded audio and transforms them into structured reports using AI-assisted workflows. The system is divided into frontend, backend, database, queueing, cache, and AI service layers.",
        ]),
        ("3.2 Requirement Analysis", [
            "The major functional requirements are user authentication and authorization, doctor and admin role support, live speech capture and file upload, structured report generation from findings, save, edit, search, and export reports, background task execution, and secure report storage. The major non-functional requirements are responsiveness, security, scalability, reliability, maintainability, and lightweight resource usage.",
        ]),
        ("3.3 Functional Modules", [
            "The system consists of authentication, speech input, transcription refinement, structured report generation, report management, search and retrieval, PDF export, and background processing modules.",
        ]),
        ("3.4 Technology Stack", [
            "The backend uses FastAPI, SQLAlchemy, PostgreSQL, Celery, and Redis. The frontend uses React, Vite, Tailwind CSS, and Quill. AI services are integrated through Hugging Face APIs. Security is handled through JWT, cookie-based session support, and AES encryption. Monitoring is supported by Flower, and migration is handled through Alembic.",
        ]),
        ("3.5 System Workflow", [
            "User opens the application and signs in. Doctor enters findings by live dictation, typed input, or audio upload. Audio is refined through API-based transcription support. Findings are passed to the structured reporting module. Study type is inferred and a suitable template is identified or cached. Structured report content is generated and formatted. The doctor reviews and edits the structured report. Final report is saved to the encrypted database. Reports can be searched, filtered, or downloaded as PDF. Heavy jobs are queued and processed in the background using Celery.",
        ]),
    ]),
    ("Chapter 4", "System Design and Architecture", [
        ("4.1 Architecture Design", [
            "The architecture of Radiology AI Platform follows a layered full-stack design with presentation, application, data, queue/cache, and AI integration layers. The presentation layer is built using React and Tailwind CSS. The application layer is built with FastAPI. PostgreSQL stores encrypted users, reports, and logs. Redis acts as broker, result backend, and caching layer. Celery processes long-running jobs asynchronously. Hugging Face APIs are used for transcript refinement and structured report assistance.",
        ]),
        ("4.2 Authentication and Security Design", [
            "Security is a major part of the platform. The authentication design includes doctor and admin roles, credential verification, token or cookie-based session management, protected routes and API authorization, and rate limiting for sensitive endpoints. The report storage design includes AES encryption for report and transcription content, controlled API access based on role, secure session handling, and Redis-backed throttling to reduce abuse.",
        ]),
        ("4.3 Speech Processing Design", [
            "The speech workflow is designed to remain lightweight. Live browser dictation provides instant UI feedback without heavy backend processing. When audio is uploaded or recording is completed, the backend receives the file and sends audio to an external speech service using Hugging Face APIs. The refined transcription is returned to the editor for user review. This avoids running heavy local ASR models and keeps the system resource-efficient.",
        ]),
        ("4.4 Structured Reporting Design", [
            "The report generation workflow is designed around structured findings rather than random free-text generation. Findings are cleaned and normalized, study type is inferred from findings, template logic is applied to define expected sections, AI support is used to populate structured JSON output, and if the AI response is weak, a deterministic fallback builds a usable structured report. The final structured report is shown in editable form on the dashboard.",
        ]),
        ("4.5 Database Design", [
            "The main database entities are User, Report, and Log. User stores user identity, email, hashed password, and role. Report stores encrypted transcription, encrypted structured report, audio hash, timestamps, and report ownership information. Log stores activity and event information for operational visibility. The database design supports report history, ownership checks, filtering, and secure storage.",
        ]),
        ("4.6 Background Task Processing", [
            "Heavy tasks such as audio-to-report processing are moved to background workers. FastAPI accepts the request, a Celery task is queued through Redis, the worker processes transcription and report generation, task status is stored and can be polled from the frontend, and results are saved and cached. This design prevents the main API from blocking while multiple doctors are using the system concurrently.",
        ]),
        ("4.7 API Design", [
            "The backend exposes authentication APIs, transcription APIs, structured report APIs, report history and save APIs, status polling APIs, and admin filtering APIs. This separation makes the application modular and easier to maintain.",
        ]),
    ]),
    ("Chapter 5", "Implementation", [
        ("5.1 Frontend Implementation", [
            "The frontend is developed using React and Vite. Tailwind CSS is used for styling and responsive layouts. The application includes a landing homepage, login and signup pages, protected report dashboard, speech input panel, editable structured report section, report history list, doctor filtering for admin users, and PDF download support.",
        ]),
        ("5.2 Backend Implementation", [
            "The backend is implemented using FastAPI and is divided into modules such as API routes, dependency injection and security, database models and session handling, AI and template services, logging and caching services, and background worker tasks. The backend has been designed with modular services including study detection, prompt building, structured reporting, report formatting, Redis caching, secure encryption, and task orchestration.",
        ]),
        ("5.3 Report Generation and Formatting", [
            "The report generation implementation includes accepting findings from the editor, running structured generation workflow, returning structured JSON and formatted text, rendering structured headings and subpoints in the UI, allowing doctors to edit the report before saving, and exporting the report as PDF with clear section hierarchy. The downloaded PDF output hides meaningless placeholder values and maintains a clean hierarchical layout for readability.",
        ]),
        ("5.4 Concurrency and Reliability Enhancements", [
            "To support multiple doctors and simultaneous requests, the following measures were implemented: Celery background workers for long-running tasks, Redis broker and result backend, connection pooling for Redis operations, retry and backoff handling for external API failures, rate limiting for high-cost endpoints, template caching and hash-based result caching, and monitoring through Flower. These enhancements make the application more practical for concurrent institutional use.",
        ]),
    ]),
    ("Chapter 6", "Results and Discussion", [
        ("6.1 Working of the System", [
            "The implemented system successfully performs the intended workflow. User authentication is handled correctly. Live dictation text appears in the editor. Uploaded audio is processed and refined. Findings are transformed into structured report output. Reports are editable before final save. Reports are stored securely and displayed in history. Reports can be searched and exported as PDF.",
        ]),
        ("6.2 Sample Outputs", [
            "The system produces refined transcription text from uploaded audio, structured report sections such as findings, impression, and recommendations, saved report entries in dashboard history, and PDF documents suitable for sharing and documentation. The generated structured report provides a much more organized format than raw free-form dictation and improves readability for downstream use.",
        ]),
        ("6.3 Testing and Validation", [
            "Major test scenarios included signup and login validation, role-based doctor or admin behavior, audio upload handling, structured report generation from typed findings, background job status tracking, save and update workflow, report retrieval and filtering, PDF export functionality, dark or light mode interface behavior, and concurrent request handling through Celery and Redis. The system behaved correctly for the intended academic project scope.",
        ]),
        ("6.4 Observations", [
            "AI-assisted structuring reduces manual report formatting effort. Background task execution improves responsiveness for heavy workflows. Encrypted storage improves confidentiality of report content. Search and PDF export increase practical utility. API-based AI usage avoids the difficulty of running large models locally. Template-driven generation improves report consistency compared to generic text output.",
        ]),
        ("6.5 Limitations", [
            "The present system still has some limitations. Output quality may depend on external AI API response quality. Network availability affects remote AI processing. Clinical verification is still required before final use of generated reports. Some specialized study types may need more domain-specific template tuning. Production deployment would require stricter infrastructure, audit, and compliance controls.",
        ]),
    ]),
    ("Chapter 7", "Conclusion", [
        ("7.1 Summary of Work", [
            "The Radiology AI Platform was developed as a full-stack AI-assisted radiology dictation and structured reporting system. It combines frontend usability, backend orchestration, secure storage, asynchronous processing, and AI-assisted formatting into a single workflow-oriented application.",
        ]),
        ("7.2 Benefits Achieved", [
            "The project reduces repetitive report-writing effort, improves structure and consistency of reports, supports better usability through a clean interface, uses scalable async design for heavy processing, maintains lightweight deployment by avoiding large local AI models, and demonstrates a practical application of AI in healthcare-oriented documentation.",
        ]),
    ]),
    ("Chapter 8", "Future Scope", [
        ("8.1 Technical Enhancements", [
            "The project can be extended in the future through more specialized templates for different radiology modalities, better domain-specific validation of generated sections, improved integration with hospital information systems and PACS, more robust audit logging and user activity analytics, better offline resilience and deployment automation, and advanced role and access management.",
        ]),
        ("8.2 Deployment and Research Opportunities", [
            "Future work may also include multi-hospital deployment support, cloud-native scaling of workers and APIs, clinical feedback loops for improving report quality, evaluation against real radiologist reporting benchmarks, multilingual dictation and reporting support, and integration with electronic health record systems.",
        ]),
    ]),
]


def add_chapters(doc: Document):
    for chapter_no, chapter_title, sections in CHAPTERS:
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_header_footer(section, chapter_title.upper())
        add_heading(doc, chapter_no, 1)
        add_heading(doc, chapter_title, 1)
        for sec_title, paras in sections:
            add_heading(doc, sec_title, 2 if sec_title.count(".") == 1 else 3)
            for para in paras:
                add_para(doc, para)


def add_references_and_appendix(doc: Document):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_header_footer(section, "REFERENCES")
    add_heading(doc, "References", 1)
    for ref in [
        "1. FastAPI Documentation. Available at: https://fastapi.tiangolo.com/",
        "2. React Documentation. Available at: https://react.dev/",
        "3. PostgreSQL Documentation. Available at: https://www.postgresql.org/docs/",
        "4. Redis Documentation. Available at: https://redis.io/docs/",
        "5. Celery Documentation. Available at: https://docs.celeryq.dev/",
        "6. Hugging Face Documentation. Available at: https://huggingface.co/docs",
        "7. Tailwind CSS Documentation. Available at: https://tailwindcss.com/docs",
        "8. jsPDF Documentation. Available at: https://github.com/parallax/jsPDF",
        "9. SQLAlchemy Documentation. Available at: https://docs.sqlalchemy.org/",
        "10. RadReport. Available at: https://radreport.org/",
        "11. Gujarat Technological University Project Report Guidelines.",
        "12. Relevant research articles on AI-assisted medical speech processing and structured radiology reporting.",
    ]:
        add_para(doc, ref)

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_header_footer(section, "APPENDIX")
    add_heading(doc, "Appendix", 1)
    for item in [
        "A. Project Title: Radiology AI Platform",
        "B. Student Details: Dev Thakar, Enrollment Number 230673107042, Computer Engineering, Sal Institute of Technology and Engineering Research, Gujarat Technological University, Academic Year 2025-26.",
        "C. Guide Details: Internal Guide - Komal Thummar, Head of Department - Dr. Nimisha Patel, Industry Mentor - Anjali Singh.",
        "D. Internship Details: FICE Education Pvt. Ltd., Remote Internship, Duration 21/01/2026 to 01/04/2026, 12 weeks.",
        "E. Core Features: doctor and admin authentication, live browser dictation, audio upload and refinement, AI-assisted structured report generation, encrypted report storage, report search and admin filtering, editable structured report interface, PDF export, async processing with Celery and Redis, monitoring through Flower.",
        "F. Suggested Figure Ideas: system architecture diagram, use case diagram, sequence diagram for report generation, database schema diagram, dashboard screenshots.",
    ]:
        add_para(doc, item)


def main():
    doc = Document()
    set_page_layout(doc)
    set_default_font(doc)
    add_cover_pages(doc)
    add_prefatory_sections(doc)
    add_lists_and_toc(doc)
    add_chapters(doc)
    add_references_and_appendix(doc)
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
